#!/usr/bin/env python3
"""
Job Processor Tool

This script reads job records from the database by runid, processes them with a local LLM
to generate CVs, cover letters, and job summaries, and stores the results in a new table.

Usage:
    python job_processor.py --runid <runid> [--database <path>] [--llm-url <url>]
"""

import argparse
import sqlite3
import json
import os
import sys
from datetime import datetime
from typing import Dict, List, Optional
import requests
from docx import Document
from docx.shared import Inches
import tempfile
import logging

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)


class JobProcessor:
    """Processes job records and generates CVs, cover letters, and summaries using a local LLM."""

    def __init__(self, database_path: str, llm_url: str = "http://localhost:11434/api/generate"):
        """
        Initialize the JobProcessor.

        Args:
            database_path: Path to the SQLite database
            llm_url: URL of the local LLM API (default: Ollama)
        """
        self.database_path = database_path
        self.llm_url = llm_url
        self.conn = None
        self.cursor = None

    def connect_database(self):
        """Connect to the SQLite database and create necessary tables."""
        try:
            self.conn = sqlite3.connect(self.database_path)
            self.cursor = self.conn.cursor()
            self._create_processed_jobs_table()
            logger.info(f"Connected to database: {self.database_path}")
        except Exception as e:
            logger.error(f"Failed to connect to database: {e}")
            raise

    def _create_processed_jobs_table(self):
        """Create the processed_jobs table if it doesn't exist."""
        self.cursor.execute('''
            CREATE TABLE IF NOT EXISTS processed_jobs (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                job_id INTEGER NOT NULL,
                runid TEXT NOT NULL,
                cv_content TEXT,
                cv_file_path TEXT,
                cover_letter TEXT,
                job_summary TEXT,
                processing_status TEXT DEFAULT 'pending',
                error_message TEXT,
                created_at DATETIME DEFAULT CURRENT_TIMESTAMP,
                updated_at DATETIME DEFAULT CURRENT_TIMESTAMP,
                FOREIGN KEY (job_id) REFERENCES jobs (id)
            )
        ''')

        # Create indexes for better performance
        self.cursor.execute('''
            CREATE INDEX IF NOT EXISTS idx_processed_jobs_runid ON processed_jobs(runid)
        ''')
        self.cursor.execute('''
            CREATE INDEX IF NOT EXISTS idx_processed_jobs_status ON processed_jobs(processing_status)
        ''')
        self.cursor.execute('''
            CREATE INDEX IF NOT EXISTS idx_processed_jobs_job_id ON processed_jobs(job_id)
        ''')

        self.conn.commit()
        logger.info("Created processed_jobs table and indexes")

    def get_new_jobs_by_runid(self, runid: str) -> List[Dict]:
        """
        Get all new job records for the given runid that haven't been processed yet.

        Args:
            runid: The run ID to filter jobs by

        Returns:
            List of job records as dictionaries
        """
        query = '''
            SELECT j.* FROM jobs j
            LEFT JOIN processed_jobs pj ON j.id = pj.job_id
            WHERE j.runid = ? AND pj.job_id IS NULL
            ORDER BY j.created_at ASC
        '''

        self.cursor.execute(query, (runid,))
        jobs = self.cursor.fetchall()

        # Get column names
        columns = [description[0] for description in self.cursor.description]

        # Convert to list of dictionaries
        job_records = []
        for job in jobs:
            job_dict = dict(zip(columns, job))
            job_records.append(job_dict)

        logger.info(f"Found {len(job_records)} new jobs for runid: {runid}")
        return job_records

    def call_local_llm(self, prompt: str, model: str = "llama3.2") -> str:
        """
        Call the local LLM API to generate content.

        Args:
            prompt: The prompt to send to the LLM
            model: The model to use (default: llama3.2)

        Returns:
            Generated text from the LLM
        """
        try:
            payload = {
                "model": model,
                "prompt": prompt,
                "stream": False
            }

            response = requests.post(self.llm_url, json=payload, timeout=120)
            response.raise_for_status()

            result = response.json()
            return result.get('response', '').strip()

        except requests.exceptions.RequestException as e:
            logger.error(f"Failed to call LLM API: {e}")
            raise
        except json.JSONDecodeError as e:
            logger.error(f"Failed to parse LLM response: {e}")
            raise

    def generate_cv_prompt(self, job_data: Dict) -> str:
        """Generate a prompt for CV creation based on job data."""
        return f"""
        Create a professional CV (resume) for the following job position. 
        The CV should be tailored to this specific role and company.
        
        Job Title: {job_data.get('job_title', 'N/A')}
        Company: {job_data.get('employer', 'N/A')}
        Location: {job_data.get('job_location', 'N/A')}
        Employment Type: {job_data.get('employment_type', 'N/A')}
        Seniority Level: {job_data.get('seniority_level', 'N/A')}
        Job Function: {job_data.get('job_function', 'N/A')}
        Industries: {job_data.get('industries', 'N/A')}
        
        Job Description:
        {job_data.get('job_description', 'N/A')}
        
        Please create a comprehensive CV that includes:
        1. Professional summary
        2. Skills section (matching the job requirements)
        3. Work experience (relevant to the position)
        4. Education
        5. Certifications (if applicable)
        6. Projects (if relevant)
        
        Make the CV professional, well-structured, and specifically tailored to this job opportunity.
        Focus on relevant skills and experience that match the job requirements.
        """

    def generate_cover_letter_prompt(self, job_data: Dict) -> str:
        """Generate a prompt for cover letter creation based on job data."""
        return f"""
        Create a compelling cover letter for the following job position.
        
        Job Title: {job_data.get('job_title', 'N/A')}
        Company: {job_data.get('employer', 'N/A')}
        Location: {job_data.get('job_location', 'N/A')}
        Employment Type: {job_data.get('employment_type', 'N/A')}
        Seniority Level: {job_data.get('seniority_level', 'N/A')}
        Job Function: {job_data.get('job_function', 'N/A')}
        Industries: {job_data.get('industries', 'N/A')}
        
        Job Description:
        {job_data.get('job_description', 'N/A')}
        
        Please create a professional cover letter that:
        1. Addresses the hiring manager professionally
        2. Explains why you're interested in this specific role and company
        3. Highlights relevant skills and experience
        4. Demonstrates understanding of the company and role
        5. Includes a strong closing statement
        6. Is concise but comprehensive (1-2 pages)
        
        Make it personal, specific to this opportunity, and compelling.
        """

    def generate_job_summary_prompt(self, job_data: Dict) -> str:
        """Generate a prompt for job summary creation based on job data."""
        return f"""
        Create a concise summary of the following job posting.
        
        Job Title: {job_data.get('job_title', 'N/A')}
        Company: {job_data.get('employer', 'N/A')}
        Location: {job_data.get('job_location', 'N/A')}
        Employment Type: {job_data.get('employment_type', 'N/A')}
        Seniority Level: {job_data.get('seniority_level', 'N/A')}
        Job Function: {job_data.get('job_function', 'N/A')}
        Industries: {job_data.get('industries', 'N/A')}
        
        Job Description:
        {job_data.get('job_description', 'N/A')}
        
        Please create a brief summary (2-3 sentences) that captures:
        1. The key responsibilities of the role
        2. The main requirements or qualifications
        3. What makes this opportunity interesting or unique
        
        Keep it concise but informative.
        """

    def create_docx_cv(self, cv_content: str, job_data: Dict) -> str:
        """
        Create a DOCX file from the CV content.

        Args:
            cv_content: The CV text content
            job_data: Job data for filename

        Returns:
            Path to the created DOCX file
        """
        try:
            # Create a new Document
            doc = Document()

            # Add title
            title = doc.add_heading(
                f"CV for {job_data.get('job_title', 'Position')} at {job_data.get('employer', 'Company')}", 0)

            # Split content into sections and add to document
            sections = cv_content.split('\n\n')
            for section in sections:
                if section.strip():
                    # Check if it's a section header (all caps or contains ':')
                    if section.isupper() or ':' in section[:50]:
                        doc.add_heading(section.strip(), level=1)
                    else:
                        doc.add_paragraph(section.strip())

            # Create filename
            safe_title = "".join(c for c in job_data.get(
                'job_title', 'job') if c.isalnum() or c in (' ', '-', '_')).rstrip()
            safe_company = "".join(c for c in job_data.get(
                'employer', 'company') if c.isalnum() or c in (' ', '-', '_')).rstrip()
            filename = f"CV_{safe_title}_{safe_company}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"

            # Create output directory if it doesn't exist
            output_dir = "generated_cvs"
            os.makedirs(output_dir, exist_ok=True)

            file_path = os.path.join(output_dir, filename)
            doc.save(file_path)

            logger.info(f"Created CV DOCX file: {file_path}")
            return file_path

        except Exception as e:
            logger.error(f"Failed to create DOCX file: {e}")
            raise

    def process_job(self, job_data: Dict) -> Dict:
        """
        Process a single job record to generate CV, cover letter, and summary.

        Args:
            job_data: Job record data

        Returns:
            Dictionary containing generated content and file paths
        """
        job_id = job_data['id']
        runid = job_data['runid']

        logger.info(
            f"Processing job {job_id}: {job_data.get('job_title', 'N/A')} at {job_data.get('employer', 'N/A')}")

        try:
            # Generate CV
            cv_prompt = self.generate_cv_prompt(job_data)
            cv_content = self.call_local_llm(cv_prompt)

            # Create DOCX file
            cv_file_path = self.create_docx_cv(cv_content, job_data)

            # Generate cover letter
            cover_letter_prompt = self.generate_cover_letter_prompt(job_data)
            cover_letter = self.call_local_llm(cover_letter_prompt)

            # Generate job summary
            summary_prompt = self.generate_job_summary_prompt(job_data)
            job_summary = self.call_local_llm(summary_prompt)

            return {
                'job_id': job_id,
                'runid': runid,
                'cv_content': cv_content,
                'cv_file_path': cv_file_path,
                'cover_letter': cover_letter,
                'job_summary': job_summary,
                'processing_status': 'completed',
                'error_message': None
            }

        except Exception as e:
            logger.error(f"Failed to process job {job_id}: {e}")
            return {
                'job_id': job_id,
                'runid': runid,
                'cv_content': None,
                'cv_file_path': None,
                'cover_letter': None,
                'job_summary': None,
                'processing_status': 'failed',
                'error_message': str(e)
            }

    def save_processed_job(self, processed_data: Dict):
        """Save processed job data to the database."""
        query = '''
            INSERT INTO processed_jobs (
                job_id, runid, cv_content, cv_file_path, cover_letter, 
                job_summary, processing_status, error_message, updated_at
            ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)
        '''

        values = (
            processed_data['job_id'],
            processed_data['runid'],
            processed_data['cv_content'],
            processed_data['cv_file_path'],
            processed_data['cover_letter'],
            processed_data['job_summary'],
            processed_data['processing_status'],
            processed_data['error_message'],
            datetime.now().isoformat()
        )

        self.cursor.execute(query, values)
        self.conn.commit()

        logger.info(
            f"Saved processed job {processed_data['job_id']} to database")

    def process_jobs_by_runid(self, runid: str):
        """
        Process all new jobs for the given runid.

        Args:
            runid: The run ID to process jobs for
        """
        logger.info(f"Starting to process jobs for runid: {runid}")

        # Get new jobs
        jobs = self.get_new_jobs_by_runid(runid)

        if not jobs:
            logger.info(f"No new jobs found for runid: {runid}")
            return

        # Process each job
        for i, job in enumerate(jobs, 1):
            logger.info(f"Processing job {i}/{len(jobs)}")

            processed_data = self.process_job(job)
            self.save_processed_job(processed_data)

            # Add a small delay between requests to be respectful to the LLM API
            import time
            time.sleep(1)

        logger.info(
            f"Completed processing {len(jobs)} jobs for runid: {runid}")

    def close(self):
        """Close the database connection."""
        if self.conn:
            self.conn.close()
            logger.info("Database connection closed")


def main():
    """Main function to run the job processor."""
    parser = argparse.ArgumentParser(
        description='Process job records and generate CVs, cover letters, and summaries')
    parser.add_argument('--runid', required=True,
                        help='Run ID to process jobs for')
    parser.add_argument('--database', default='webapp/jobs.db',
                        help='Path to the SQLite database (default: webapp/jobs.db)')
    parser.add_argument('--llm-url', default='http://localhost:11434/api/generate',
                        help='URL of the local LLM API (default: http://localhost:11434/api/generate)')
    parser.add_argument('--model', default='llama3.2',
                        help='LLM model to use (default: llama3.2)')

    args = parser.parse_args()

    # Check if database exists
    if not os.path.exists(args.database):
        logger.error(f"Database file not found: {args.database}")
        sys.exit(1)

    # Initialize processor
    processor = JobProcessor(args.database, args.llm_url)

    try:
        # Connect to database
        processor.connect_database()

        # Process jobs
        processor.process_jobs_by_runid(args.runid)

        logger.info("Job processing completed successfully")

    except Exception as e:
        logger.error(f"Job processing failed: {e}")
        sys.exit(1)
    finally:
        processor.close()


if __name__ == "__main__":
    main()
